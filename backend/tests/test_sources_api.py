"""Sources API 与定稿刷新集成测试。"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document


async def _create_project(auth_client, title: str = "Sources Proj") -> str:
    res = await auth_client.post("/api/projects", json={"title": title})
    assert res.status_code == 201
    return res.json()["id"]


async def test_paste_assessment_refreshes_project_summary(auth_client):
    pid = await _create_project(auth_client)
    create = await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={
            "role": "ASSESSMENT",
            "source_type": "PASTE",
            "title": "Rubric",
            "raw_text": "Write an APA 7th literature review on AI tutoring.",
        },
    )
    assert create.status_code == 201
    body = create.json()
    assert body["role"] == "ASSESSMENT"
    assert body["status"] == "SUMMARIZED"
    assert "APA 7th" in (body["summary_text"] or "")

    project = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert project["assessment_ready"] is True
    assert "AI tutoring" in project["assessment_summary"]
    assert project["source_document_count"] == 1


async def test_list_sources_filter_by_role(auth_client):
    pid = await _create_project(auth_client, "Filter")
    await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={"role": "ASSESSMENT", "raw_text": "A text"},
    )
    await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={"role": "BACKGROUND", "raw_text": "B text"},
    )
    all_docs = await auth_client.get(f"/api/projects/{pid}/sources")
    assert all_docs.status_code == 200
    assert len(all_docs.json()) == 2

    only_b = await auth_client.get(f"/api/projects/{pid}/sources?role=BACKGROUND")
    assert only_b.status_code == 200
    assert len(only_b.json()) == 1
    assert only_b.json()[0]["role"] == "BACKGROUND"


async def test_specific_refresh_and_delete(auth_client):
    pid = await _create_project(auth_client, "Specific")
    created = await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={"role": "SPECIFIC", "raw_text": "Max 3000 words. Use APA."},
    )
    sid = created.json()["id"]
    project = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert "3000 words" in project["specific_requirements"]

    deleted = await auth_client.delete(f"/api/projects/{pid}/sources/{sid}")
    assert deleted.status_code == 204
    project = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert project["specific_requirements"] is None
    assert project["source_document_count"] == 0


async def test_background_does_not_touch_assembled_fields(auth_client):
    pid = await _create_project(auth_client, "BgOnly")
    await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={"role": "ASSESSMENT", "raw_text": "Keep assessment"},
    )
    before = (await auth_client.get(f"/api/projects/{pid}")).json()
    await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={"role": "BACKGROUND", "raw_text": "Notebook notes should not overwrite A"},
    )
    after = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert after["assessment_summary"] == before["assessment_summary"]
    assert after["paper_outline"] is None
    assert after["source_document_count"] == 2


async def test_outline_lock_sets_paper_outline(auth_client):
    pid = await _create_project(auth_client, "Outline")
    md = "# Introduction\n\nScope here.\n\n## Methods\n\nApproach.\n\n# Conclusion\n"
    created = await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={"role": "OUTLINE", "title": "Structure", "raw_text": md},
    )
    assert created.status_code == 201
    doc = created.json()
    assert doc["status"] == "SUMMARIZED"
    assert isinstance(doc["summary_json"], list)
    assert len(doc["summary_json"]) >= 2

    # 未锁定前 outline_ready 为 false
    project = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert project["outline_ready"] is False

    locked = await auth_client.post(
        f"/api/projects/{pid}/outline/lock",
        json={"source_id": doc["id"]},
    )
    assert locked.status_code == 200

    project = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert project["outline_ready"] is True
    assert project["outline_locked_at"]
    headings = [x["heading"] for x in project["paper_outline"]]
    assert "Introduction" in headings
    assert "Methods" in headings


async def test_upload_markdown_source(auth_client):
    pid = await _create_project(auth_client, "UploadMd")
    content = b"# Rubric\n\nCriterion: critical analysis.\n"
    files = {"file": ("rubric.md", BytesIO(content), "text/markdown")}
    data = {"role": "ASSESSMENT", "title": "Uploaded rubric"}
    res = await auth_client.post(
        f"/api/projects/{pid}/sources/upload",
        data=data,
        files=files,
    )
    assert res.status_code == 201
    body = res.json()
    assert body["source_type"] == "UPLOAD"
    assert body["original_filename"] == "rubric.md"
    assert body["storage_path"]
    assert Path(body["storage_path"]).exists()
    assert "critical analysis" in (body["raw_text"] or "")

    project = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert project["assessment_ready"] is True


async def test_upload_docx_outline_and_lock(auth_client):
    pid = await _create_project(auth_client, "UploadDocx")
    buf = BytesIO()
    document = Document()
    document.add_heading("Literature Review", level=1)
    document.add_paragraph("Overview of the field.")
    document.add_heading("Findings", level=2)
    document.add_paragraph("Key results.")
    document.save(buf)
    buf.seek(0)

    res = await auth_client.post(
        f"/api/projects/{pid}/sources/upload",
        data={"role": "OUTLINE"},
        files={
            "file": (
                "outline.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 201
    sid = res.json()["id"]
    assert res.json()["status"] == "SUMMARIZED"

    lock = await auth_client.post(
        f"/api/projects/{pid}/outline/lock",
        json={"source_id": sid},
    )
    assert lock.status_code == 200
    project = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert project["outline_ready"] is True
    assert any(x["heading"] == "Literature Review" for x in project["paper_outline"])


async def test_reparse_source(auth_client):
    pid = await _create_project(auth_client, "Reparse")
    created = await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={"role": "OUTLINE", "raw_text": "# Only One"},
    )
    sid = created.json()["id"]
    # 直接改库不方便；reparse 后应重新摘要为 SUMMARIZED
    reparsed = await auth_client.post(f"/api/projects/{pid}/sources/{sid}/reparse")
    assert reparsed.status_code == 200
    assert reparsed.json()["status"] == "SUMMARIZED"
    assert reparsed.json()["summary_json"][0]["heading"] == "Only One"


async def test_notebook_sync_creates_background(auth_client, monkeypatch):
    pid = await _create_project(auth_client, "NbSync")
    url = "https://notebooklm.google.com/notebook/demo"

    async def _fake_fetch(notebook_url: str) -> str:
        assert notebook_url == url
        return "User:\nHello\nNotebookLM:\nWorld about education research."

    from app.services import notebooklm as notebooklm_module

    monkeypatch.setattr(
        notebooklm_module.notebooklm_service,
        "fetch_via_browser",
        _fake_fetch,
    )

    res = await auth_client.post(
        f"/api/projects/{pid}/sources/notebook-sync",
        json={"notebook_url": url, "use_browser": True},
    )
    assert res.status_code == 201
    body = res.json()
    assert body["role"] == "BACKGROUND"
    assert body["source_type"] == "NOTEBOOKLM"
    assert body["notebook_url"] == url
    assert "education research" in (body["raw_text"] or "")

    # B 不写定稿
    project = (await auth_client.get(f"/api/projects/{pid}")).json()
    assert project["assessment_summary"] is None
    assert project["outline_ready"] is False


async def test_invalid_role_rejected(auth_client):
    pid = await _create_project(auth_client, "BadRole")
    res = await auth_client.post(
        f"/api/projects/{pid}/sources",
        json={"role": "WRONG", "raw_text": "x"},
    )
    assert res.status_code == 400


async def test_old_notebook_route_still_gone(auth_client):
    pid = await _create_project(auth_client, "NoOldNb")
    res = await auth_client.get(f"/api/notebook/{pid}")
    assert res.status_code == 404
