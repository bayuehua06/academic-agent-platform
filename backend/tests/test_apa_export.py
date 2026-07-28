"""APA DOCX 导出与文件名测试。"""

from pathlib import Path

from docx import Document

from app.services.apa_docx import (
    build_apa_reference_docx,
    markdown_to_apa_docx,
    sanitize_export_stem,
)
from app.services.pandoc_service import pandoc_service


def test_sanitize_export_stem():
    assert sanitize_export_stem("My Paper / Draft", 3) == "My_Paper_Draft_v3"
    assert sanitize_export_stem("My Paper", "1.1") == "My_Paper_v1.1"
    assert sanitize_export_stem("  ", 1) == "draft_v1"
    assert ":" not in sanitize_export_stem("a:b?", 2)


def test_markdown_to_apa_docx_styles(tmp_path: Path):
    md = """## Introduction

This is a body paragraph with *italic* and **bold** text.

## References

Smith, J. (2024). *Sample title*. Journal, 1(2), 3-4.
"""
    out = tmp_path / "sample.docx"
    markdown_to_apa_docx(md, out)
    doc = Document(str(out))
    section = doc.sections[0]
    assert abs(section.top_margin.inches - 1.0) < 0.05
    headings = [p.text for p in doc.paragraphs if p.style and "Heading" in (p.style.name or "")]
    assert "Introduction" in headings
    assert "References" in headings
    # References 条目应有悬挂缩进
    ref_paras = [p for p in doc.paragraphs if "Smith" in p.text]
    assert ref_paras
    assert ref_paras[0].paragraph_format.left_indent is not None


def test_markdown_to_apa_docx_tables(tmp_path: Path):
    md = """## Methods

Intro sentence.

| Theme | Focus |
| --- | --- |
| Platformization | Gig economy |
| Labor | Control |

After the table.
"""
    out = tmp_path / "with_table.docx"
    markdown_to_apa_docx(md, out)
    doc = Document(str(out))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text.strip() == "Theme"
    assert "Platformization" in table.rows[1].cells[0].text
    assert "Gig economy" in table.rows[1].cells[1].text
    # 不应把 | 管道符当正文残留
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "| Theme |" not in body
    assert "After the table." in body


async def test_export_filename_content_disposition(auth_client, monkeypatch):
    from unittest.mock import patch

    from app.services import summarizer as summarizer_module
    from tests.helpers import prepare_confirmed_literatures, prepare_writing_inputs

    monkeypatch.setattr(summarizer_module, "has_openai_key", lambda: False)

    create = await auth_client.post(
        "/api/projects",
        json={"title": "Food Delivery Review"},
    )
    pid = create.json()["id"]
    await prepare_writing_inputs(auth_client, pid)
    _lits, mock_svc = await prepare_confirmed_literatures(auth_client, pid, count=1)
    with patch("app.services.literature_workflow.zotero_for_project", return_value=mock_svc):
        run = await auth_client.post(
            f"/api/projects/{pid}/run-agent",
            json={"max_papers": 1, "skip_search": True},
        )
    assert run.status_code == 200, run.text
    res = await auth_client.get(f"/api/drafts/{pid}/export?format=docx")
    assert res.status_code == 200
    cd = res.headers.get("content-disposition", "")
    assert "Food_Delivery_Review_v1.docx" in cd


def test_build_reference_docx(tmp_path: Path):
    path = tmp_path / "apa_reference.docx"
    build_apa_reference_docx(path)
    assert path.exists()
    assert path.stat().st_size > 1000


def test_pandoc_service_docx_apa(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(pandoc_service, "export_dir", tmp_path)
    path = pandoc_service.markdown_to_document(
        "## Methods\n\nBody text here.\n",
        "Doe, A. (2020). Title.",
        fmt="docx",
        filename_stem="Demo_Paper_v2",
    )
    assert path.name == "Demo_Paper_v2.docx"
    doc = Document(str(path))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Methods" in texts
    assert "Doe" in texts


def test_merge_replaces_truncated_writer_references():
    from app.services.pandoc_service import merge_content_with_apa_references

    body = (
        "## Intro\n\nHello.\n\n## References\n\n"
        "Kostadinova-Tsankova, L., Stoyanov, S., Tabak"
    )
    full = "Smith, J. (2024). *Complete paper*. Journal.\nhttps://doi.org/10.1/xyz"
    merged = merge_content_with_apa_references(body, full)
    assert "Tabak" not in merged
    assert "Complete paper" in merged
    assert merged.count("## References") == 1
    assert "Intro" in merged
