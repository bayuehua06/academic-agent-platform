"""APA 7th 学生论文风格的 Markdown → DOCX（python-docx）。

在无 pandoc / 无 reference.docx 时作为主路径；有 pandoc 时也可作为回退。
版式要点：Times New Roman 12、双倍行距、1 英寸边距、标题层级、正文首行缩进、References 悬挂缩进。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


_EMPHASIS_RE = re.compile(
    r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_)",
    re.DOTALL,
)


def sanitize_export_stem(title: str, version_label: int | str) -> str:
    """项目名 + 版本标签 → 安全文件名主干（不含扩展名）。"""
    raw = (title or "").strip() or "draft"
    cleaned = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", raw)
    cleaned = re.sub(r"\s+", "_", cleaned)
    cleaned = re.sub(r"_+", "_", cleaned).strip("._")
    if len(cleaned) > 80:
        cleaned = cleaned[:80].rstrip("._")
    if not cleaned:
        cleaned = "draft"
    label = str(version_label).strip() or "1"
    label = re.sub(r'[\\/:*?"<>|\s]+', "_", label)
    return f"{cleaned}_v{label}"


def _set_run_font(run, *, bold: bool = False, italic: bool = False, size_pt: float = 12) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 东亚字体回退，避免部分环境下中文乱码样式
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _set_paragraph_double_spacing(paragraph, *, first_line_inches: Optional[float] = 0.5) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_inches is None:
        pf.first_line_indent = None
    else:
        pf.first_line_indent = Inches(first_line_inches)


def _add_page_number(paragraph) -> None:
    """在段落中插入 PAGE 域。"""
    run = paragraph.add_run()
    _set_run_font(run, size_pt=12)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_char_begin)
    r_element.append(instr_text)
    r_element.append(fld_char_separate)
    r_element.append(fld_char_end)


def configure_apa_document(doc: Document) -> None:
    """设置页面边距、页码与 Normal / Heading 样式。"""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # 清空再加页码
        for run in list(hp.runs):
            run.text = ""
        _add_page_number(hp)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    nf = normal.paragraph_format
    nf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    nf.space_before = Pt(0)
    nf.space_after = Pt(0)
    nf.first_line_indent = Inches(0.5)

    # Heading 1–3：APA 层级（字号仍 12）
    for level, align, bold, italic in (
        (1, WD_ALIGN_PARAGRAPH.CENTER, True, False),
        (2, WD_ALIGN_PARAGRAPH.LEFT, True, False),
        (3, WD_ALIGN_PARAGRAPH.LEFT, True, True),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.alignment = align
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Inches(0)


def _add_runs_with_emphasis(paragraph, text: str) -> None:
    """解析简单 *斜体* / **粗体** / ***粗斜体***。"""
    pos = 0
    for match in _EMPHASIS_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            _set_run_font(run)
        if match.group(2) is not None:
            run = paragraph.add_run(match.group(2))
            _set_run_font(run, bold=True, italic=True)
        elif match.group(3) is not None:
            run = paragraph.add_run(match.group(3))
            _set_run_font(run, bold=True)
        elif match.group(4) is not None or match.group(5) is not None:
            run = paragraph.add_run(match.group(4) or match.group(5))
            _set_run_font(run, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run)
    if not paragraph.runs:
        run = paragraph.add_run(text)
        _set_run_font(run)


def _strip_md_link(text: str) -> str:
    """[label](url) → label；保留裸 URL。"""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("\\", "")


def _is_references_heading(text: str) -> bool:
    t = text.strip().lower().rstrip(":")
    return t in {"references", "reference", "参考文献"}


def _is_md_table_row(line: str) -> bool:
    """是否像 Markdown 表格行（含分隔行）。"""
    s = line.strip()
    if "|" not in s:
        return False
    # 至少两段（一个 | 不够稳；要求两侧或中间有单元格）
    parts = [p.strip() for p in s.strip("|").split("|")]
    return len(parts) >= 2


def _is_md_table_separator(line: str) -> bool:
    """| --- | :---: | ---: |"""
    s = line.strip().strip("|")
    if not s or "|" not in line:
        return False
    cells = [c.strip() for c in s.split("|")]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells if c)


def _parse_md_table_row(line: str) -> List[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip().replace("\\|", "|") for c in s.split("|")]


def _add_word_table(doc: Document, rows: List[List[str]]) -> None:
    """把二维单元格写成 Word 表格（首行加粗作表头）。"""
    if not rows:
        return
    width = max(len(r) for r in rows)
    if width < 1:
        return
    norm = [r + [""] * (width - len(r)) for r in rows]
    table = doc.add_table(rows=len(norm), cols=width)
    table.style = "Table Grid"
    for i, row_cells in enumerate(norm):
        for j, cell_text in enumerate(row_cells):
            cell = table.rows[i].cells[j]
            # 清空默认空段再写
            cell.text = ""
            p = cell.paragraphs[0]
            _set_paragraph_double_spacing(p, first_line_inches=0)
            plain = _strip_md_link(cell_text)
            plain = re.sub(r"[*_`]+", "", plain).strip() if i == 0 else plain
            if i == 0:
                run = p.add_run(plain)
                _set_run_font(run, bold=True)
            else:
                _add_runs_with_emphasis(p, plain)
    # 表后空一行间距（用空段）
    spacer = doc.add_paragraph()
    _set_paragraph_double_spacing(spacer, first_line_inches=0)


def markdown_to_apa_docx(markdown_text: str, output_path: Path) -> Path:
    """将 Markdown 草稿写成 APA 风格 docx。"""
    doc = Document()
    configure_apa_document(doc)

    in_references = False
    buffer: List[str] = []
    lines = markdown_text.splitlines()
    idx = 0

    def flush_paragraph() -> None:
        nonlocal buffer
        if not buffer:
            return
        text = _strip_md_link(" ".join(buffer).strip())
        buffer = []
        if not text:
            return
        p = doc.add_paragraph()
        if in_references:
            _set_paragraph_double_spacing(p, first_line_inches=-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
        else:
            _set_paragraph_double_spacing(p, first_line_inches=0.5)
        _add_runs_with_emphasis(p, text)

    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            idx += 1
            continue

        # Markdown 表格：连续 | 行（跳过 --- 分隔行）
        if _is_md_table_row(stripped) and not in_references:
            # 需要至少表头 + 分隔，或 ≥2 行数据，才当表；单行带 | 仍当正文
            look_ahead = [stripped]
            j = idx + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    break
                if _is_md_table_row(nxt):
                    look_ahead.append(nxt)
                    j += 1
                    continue
                break
            data_rows = [_parse_md_table_row(r) for r in look_ahead if not _is_md_table_separator(r)]
            has_sep = any(_is_md_table_separator(r) for r in look_ahead)
            if has_sep or len(data_rows) >= 2:
                flush_paragraph()
                _add_word_table(doc, data_rows)
                idx = j
                continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            title = _strip_md_link(heading_match.group(2).strip())
            # 去掉 Markdown 标题里的强调标记再加样式
            title_plain = re.sub(r"[*_`]+", "", title).strip()
            if _is_references_heading(title_plain):
                in_references = True
            p = doc.add_paragraph(style=f"Heading {level}")
            # Heading 样式已设对齐；再写 run 保证字体
            for run in list(p.runs):
                run.text = ""
            run = p.add_run(title_plain)
            _set_run_font(
                run,
                bold=True,
                italic=(level == 3),
                size_pt=12,
            )
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_double_spacing(p, first_line_inches=0)
            idx += 1
            continue

        # 无序列表 → 普通段落（APA 作业稿少用 bullet）
        if re.match(r"^[-*+]\s+", stripped):
            stripped = re.sub(r"^[-*+]\s+", "", stripped)
        elif re.match(r"^\d+\.\s+", stripped):
            stripped = re.sub(r"^\d+\.\s+", "", stripped)

        buffer.append(stripped)
        idx += 1

    flush_paragraph()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def build_apa_reference_docx(path: Path) -> Path:
    """生成供 pandoc `--reference-doc` 使用的空白 APA 模板。"""
    doc = Document()
    configure_apa_document(doc)
    # 放一段示例 Normal，确保样式写入
    p = doc.add_paragraph("Sample body text for APA reference styles.")
    _set_paragraph_double_spacing(p, first_line_inches=0.5)
    for run in p.runs:
        _set_run_font(run)
    for level in (1, 2, 3):
        hp = doc.add_paragraph(f"Heading {level} Sample", style=f"Heading {level}")
        _set_paragraph_double_spacing(hp, first_line_inches=0)
        for run in hp.runs:
            _set_run_font(run, bold=True, italic=(level == 3))
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
