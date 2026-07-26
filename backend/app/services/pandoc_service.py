"""Pandoc / python-docx 文档转换服务。"""

from __future__ import annotations

import logging
import shutil
from pathlib import Path
from typing import Literal, Optional
from uuid import uuid4

from app.core.config import get_settings
from app.services.apa_docx import build_apa_reference_docx, markdown_to_apa_docx

logger = logging.getLogger(__name__)
settings = get_settings()

ExportFormat = Literal["docx", "pdf"]


class PandocService:
    """Markdown ↔ Word/PDF 转换引擎。"""

    def __init__(
        self,
        export_dir: Optional[str] = None,
        apa_csl_path: Optional[str] = None,
        apa_reference_docx: Optional[str] = None,
    ) -> None:
        self.export_dir = Path(export_dir or settings.export_dir)
        self.apa_csl_path = Path(apa_csl_path or settings.apa_csl_path)
        self.apa_reference_docx = Path(
            apa_reference_docx or getattr(settings, "apa_reference_docx", "./resources/apa_reference.docx")
        )
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def _ensure_reference_docx(self) -> Path:
        """确保 APA reference.docx 存在（供 pandoc 使用）。"""
        if not self.apa_reference_docx.exists():
            build_apa_reference_docx(self.apa_reference_docx)
        return self.apa_reference_docx

    def markdown_to_document(
        self,
        content_markdown: str,
        apa_references_block: Optional[str] = None,
        fmt: ExportFormat = "docx",
        filename_stem: Optional[str] = None,
        *,
        prefer_apa_python: bool = True,
    ) -> Path:
        """
        将 Markdown 正文 + APA References 导出为 docx 或 pdf。

        docx 默认走 python-docx APA 版式（不依赖本机 pandoc）；
        若 prefer_apa_python=False 且已安装 pandoc，则用 pandoc + reference-doc。
        """
        full_md = merge_content_with_apa_references(content_markdown, apa_references_block)

        stem = filename_stem or f"draft_{uuid4().hex[:8]}"
        output_path = self.export_dir / f"{stem}.{fmt}"

        if fmt == "docx" and prefer_apa_python:
            return markdown_to_apa_docx(full_md, output_path)

        if shutil.which("pandoc"):
            try:
                import pypandoc

                extra_args = ["--standalone"]
                if fmt == "docx":
                    ref = self._ensure_reference_docx()
                    extra_args.extend(["--reference-doc", str(ref)])
                    if self.apa_csl_path.exists():
                        extra_args.extend(["--csl", str(self.apa_csl_path)])

                pypandoc.convert_text(
                    full_md,
                    to=fmt,
                    format="markdown",
                    outputfile=str(output_path),
                    extra_args=extra_args,
                )
                return output_path
            except Exception as exc:  # noqa: BLE001
                logger.warning("pypandoc 转换失败，回退 APA python-docx: %s", exc)

        if fmt == "pdf":
            raise RuntimeError("PDF 导出需要系统安装 pandoc（及可选 pdflatex/weasyprint）")

        return markdown_to_apa_docx(full_md, output_path)

    def docx_to_markdown(self, docx_path: Path) -> str:
        """将 Word 文档转为 Markdown。"""
        if shutil.which("pandoc"):
            try:
                import pypandoc

                return pypandoc.convert_file(str(docx_path), to="md", format="docx")
            except Exception as exc:  # noqa: BLE001
                logger.warning("pypandoc 导入失败，回退 python-docx: %s", exc)

        return self._fallback_docx_to_md(docx_path)

    def _fallback_docx_to_md(self, docx_path: Path) -> str:
        """无 pandoc 时用 python-docx 提取段落与表格为 Markdown。"""
        from docx import Document

        doc = Document(str(docx_path))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if "heading 1" in style:
                lines.append(f"# {text}")
            elif "heading 2" in style:
                lines.append(f"## {text}")
            elif "heading 3" in style:
                lines.append(f"### {text}")
            else:
                lines.append(text)

        for table_idx, table in enumerate(doc.tables):
            if table_idx > 0 or lines:
                lines.append("")
            lines.append(f"### Table {table_idx + 1}")
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = " ".join(
                        p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()
                    )
                    cells.append(cell_text.replace("|", "\\|") if cell_text else "")
                if not any(cells):
                    continue
                lines.append("| " + " | ".join(cells) + " |")

        return "\n\n".join(line for line in lines if line is not None).strip() + "\n"

    def save_upload(self, content: bytes, suffix: str = ".docx") -> Path:
        """保存上传文件到临时目录。"""
        upload_root = Path(settings.upload_dir)
        upload_root.mkdir(parents=True, exist_ok=True)
        path = upload_root / f"{uuid4().hex}{suffix}"
        path.write_bytes(content)
        return path


def re_has_references_heading(markdown: str) -> bool:
    """正文是否已含 References 标题。"""
    import re

    return bool(
        re.search(r"(?im)^#{1,3}\s+references?\s*$", markdown)
        or re.search(r"(?im)^#{1,3}\s+参考文献\s*$", markdown)
    )


def strip_references_section(markdown: str) -> str:
    """
    去掉正文末尾（或中部起）的 References / 参考文献章节。

    Writer 常在 token 截断时留下半截 References；导出应以 apa_references_block 为准。
    """
    import re

    text = (markdown or "").rstrip()
    if not text:
        return ""
    match = re.search(r"(?im)^#{1,3}\s+(references?|参考文献)\s*$", text)
    if not match:
        return text
    return text[: match.start()].rstrip()


def merge_content_with_apa_references(
    content_markdown: str,
    apa_references_block: Optional[str],
) -> str:
    """合并正文与权威 References；有 block 时总是替换正文里的 References 段。"""
    body = strip_references_section(content_markdown).strip()
    refs = (apa_references_block or "").strip()
    if refs:
        return f"{body}\n\n## References\n\n{refs}\n" if body else f"## References\n\n{refs}\n"
    # 无权威 block 时保留原文（含可能不完整的 References）
    return (content_markdown or "").strip() + ("\n" if content_markdown else "")


pandoc_service = PandocService()
